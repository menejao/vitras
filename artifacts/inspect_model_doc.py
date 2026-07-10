# -*- coding: utf-8 -*-
from docx import Document
path = r"C:\Users\joão.benedito\Downloads\Escopo Funcional -34966- Importador NFe - Unificado novo V1.2.docx"
doc = Document(path)
print('paragraphs', len(doc.paragraphs))
count = 0
for i,p in enumerate(doc.paragraphs):
    txt = p.text.strip().replace('\t',' ')
    if not txt:
        continue
    print(f"P{i:03d}|{p.style.name}|{txt[:180]}")
    count += 1
    if count >= 160:
        break
print('tables', len(doc.tables))
for ti,t in enumerate(doc.tables[:20]):
    rows=len(t.rows); cols=len(t.columns)
    print(f"T{ti:02d}|{rows}x{cols}")
    for r in t.rows[:5]:
        vals=[c.text.strip().replace('\n',' / ')[:80] for c in r.cells]
        print('  ', vals)
