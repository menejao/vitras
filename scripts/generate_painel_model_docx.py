from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\dev\vitras")
SOURCE_MD = ROOT / "escopos" / "painel.md"
OUTPUT_DIR = ROOT / "escopos"
OUTPUT_DOCX = OUTPUT_DIR / "painel.docx"

BLUE = RGBColor(31, 73, 125)
TEXT = RGBColor(34, 34, 34)
FILL_HEADER = "D9E2F3"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, color=TEXT, size=10, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths):
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, FILL_HEADER)
            set_repeat_table_header(row)


def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(4)


def add_line(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT


def add_table_generic(doc, heading, headers, rows, widths):
    add_line(doc, heading)
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, color=BLUE, align="center")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def build_doc():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    add_title(doc, "Escopo Funcional - Painel")

    add_table_generic(doc, "Identificação do documento:", ["Campo", "Descrição"], [
        ("Sistema", "VITRAS APS"),
        ("Tela", "Painel"),
        ("Objetivo", "Visão operacional resumida da unidade"),
        ("Versão", "1.0"),
        ("Arquivo base", "escopos/painel.md"),
    ], [Cm(4.5), Cm(12.0)])

    add_line(doc, "Objetivo e contexto:")
    add_line(doc, "Permitir leitura rápida da situação operacional da unidade logo na entrada do sistema, com indicadores, alertas e atalhos para aprofundamento.")
    add_line(doc, "Localização da tela: menu principal > Painel.")

    add_table_generic(doc, "Informações principais da tela:", ["Bloco", "Função"], [
        ("KPIs", "Resumo rápido do cenário atual da unidade"),
        ("Demanda programada", "Leitura do equilíbrio entre demanda programada e espontânea"),
        ("Alertas proativos", "Destaque para situações que exigem atenção"),
        ("Pacientes prioritários", "Atalho para casos mais sensíveis"),
        ("Equipe ativa", "Resumo visual dos profissionais considerados na tela"),
    ], [Cm(5.0), Cm(11.5)])

    add_table_generic(doc, "Ações:", ["Ação", "Comportamento esperado"], [
        ("Abrir gestão à vista", "Encaminhar para visão mais gerencial"),
        ("Ir para pacientes", "Encaminhar para lista de pacientes"),
        ("Clique em alerta", "Quando existir paciente vinculado, abrir seu contexto"),
        ("Clique em paciente prioritário", "Abrir contexto do paciente correspondente"),
    ], [Cm(5.0), Cm(11.5)])

    add_table_generic(doc, "Dicionário de campos:", ["Campo", "Função", "Como aparece internamente", "Observações"], [
        ("Painel", "Nome funcional da tela", "dashboard", "Nome de navegação interna"),
        ("Pacientes ativos", "Mostrar total de pacientes", "patients.length", "Indicador de volume"),
        ("Com ACS definido", "Mostrar pacientes com ACS vinculado", "assignedAcsId", "Indicador de cobertura"),
        ("Protocolos críticos", "Mostrar situações mais sensíveis", "protocolByPatient / protocolChip", "Indicador de prioridade"),
        ("Profissionais", "Mostrar total de profissionais", "users.length", "Resumo da equipe"),
        ("Demanda programada", "Mostrar percentual programado", "demandMonthly", "Pode ficar sem dados"),
        ("Alertas proativos", "Destacar situações de atenção", "alerts", "Lista resumida"),
        ("Pacientes prioritários", "Listar pacientes críticos", "critical", "Lista curta"),
    ], [Cm(3.5), Cm(5.0), Cm(4.0), Cm(4.0)])

    add_table_generic(doc, "Permissões de acesso:", ["Perfil", "Comportamento esperado"], [
        ("Perfis assistenciais", "Acessam o Painel como visão operacional inicial"),
        ("Gestor", "Pode acessar, mas tende a seguir para visão mais gerencial"),
        ("Recepção", "Usa fluxo próprio e não depende do Painel principal"),
        ("Suporte e auditoria", "Acesso conforme política da plataforma"),
    ], [Cm(4.5), Cm(12.0)])

    add_table_generic(doc, "Auditoria das interações:", ["Evento", "Rastreabilidade esperada"], [
        ("Leitura inicial do painel", "Registrar abertura da visão operacional"),
        ("Uso de atalho do cabeçalho", "Registrar navegação iniciada no Painel"),
        ("Abertura por alerta", "Registrar aprofundamento motivado por alerta"),
        ("Abertura por paciente prioritário", "Registrar aprofundamento por priorização"),
    ], [Cm(5.5), Cm(11.0)])

    add_table_generic(doc, "Critérios de aceitação de testes", ["Dado", "Quando", "Então"], [
        ("Usuário com acesso ao Painel", "Abrir sistema", "Visualiza panorama operacional resumido da unidade"),
        ("Dados de pacientes e protocolos disponíveis", "Tela carregar", "KPIs refletem o cenário carregado"),
        ("Existirem pacientes críticos", "Tela montar lista prioritária", "Pacientes prioritários aparecem com destaque"),
        ("Existirem alertas relevantes", "Tela montar alertas", "Alertas aparecem de forma resumida e acionável"),
        ("Não houver base para demanda programada", "Tela carregar bloco de demanda", "Sistema mostra 'Sem dados'"),
    ], [Cm(5.5), Cm(4.5), Cm(7.0)])

    add_table_generic(doc, "Cenários de testes", ["Cenário", "Resultado esperado"], [
        ("Abertura do Painel com dados completos", "Tela carrega panorama geral sem inconsistência"),
        ("Abertura do Painel sem pacientes críticos", "Mensagem de vazio aparece no bloco correspondente"),
        ("Abertura do Painel sem alertas", "Seção não aparenta falha"),
        ("Abertura com demanda programada disponível", "Percentual e status aparecem corretamente"),
        ("Clique em paciente prioritário", "Sistema leva ao contexto do paciente"),
        ("Fluxo de recepção", "Usuário não depende do Painel principal"),
    ], [Cm(6.0), Cm(10.5)])

    add_line(doc, "Pontos de atenção:")
    add_line(doc, "Existem sinais de desalinhamento entre texto de algumas ações rápidas e comportamento efetivamente realizado pela navegação.")
    add_line(doc, "Há oportunidade de tornar o Painel mais útil com melhor aproveitamento de qualidade de dados e melhor ligação entre alertas e ação prática.")

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_doc()
