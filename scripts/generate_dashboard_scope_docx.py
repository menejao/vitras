from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\dev\vitras")
OUTPUT = ROOT / "docs" / "escopos" / "painel.docx"

ACCENT_DARK = RGBColor(24, 45, 74)
TEXT = RGBColor(40, 40, 40)
MUTED = RGBColor(92, 104, 117)
HEADER_FILL = "D7E6F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=TEXT, size=10.5, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, widths=None):
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "left", "bottom", "right"):
                element = tc_mar.find(qn(f"w:{side}"))
                if element is None:
                    element = OxmlElement(f"w:{side}")
                    tc_mar.append(element)
                element.set(qn("w:w"), "110")
                element.set(qn("w:type"), "dxa")
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, HEADER_FILL)
            set_repeat_table_header(row)


def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT_DARK
    p.space_after = Pt(6)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Aptos"
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    p.space_after = Pt(12)


def add_meta_table(doc):
    table = doc.add_table(rows=4, cols=2)
    style_table(table, [Cm(4.2), Cm(11.8)])
    rows = [
        ("Produto", "VITRAS"),
        ("Módulo", "Dashboard / Visão Operacional da Unidade"),
        ("Tab/Rota", "`dashboard`"),
        ("Última atualização", "27/05/2026"),
    ]
    for i, (label, value) in enumerate(rows):
        set_cell_text(table.cell(i, 0), label, bold=True, color=ACCENT_DARK)
        set_cell_text(table.cell(i, 1), value)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    r.font.color.rgb = ACCENT_DARK
    r.font.name = "Aptos Display"
    r.font.size = Pt(15 if level == 1 else 12.5 if level == 2 else 11.5)
    p.space_before = Pt(8)
    p.space_after = Pt(4)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT
        p.paragraph_format.space_after = Pt(1)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT
        p.paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, color=ACCENT_DARK, align="center")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_rule(doc, rid, desc, impact, exc):
    for prefix, text in [
        (f"{rid} — Descrição: ", desc),
        ("Impacto: ", impact),
        ("Exceção: ", exc),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(prefix)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = ACCENT_DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = TEXT


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(1.9)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "Painel")
    add_subtitle(doc, "Especificação funcional enterprise da visão operacional da unidade.")
    add_meta_table(doc)

    add_heading(doc, "1. Objetivo da Tela")
    add_heading(doc, "Qual problema resolve", level=2)
    add_bullets(doc, [
        "Consolidar visão operacional da UBS em uma única tela de entrada.",
        "Reduzir navegação entre módulos para identificar pacientes críticos, alertas e estado geral da equipe.",
        "Transformar dados dispersos de pacientes, protocolos, agenda e farmácia em briefing operacional rápido.",
        "Ajudar coordenação clínica a agir antes que atraso de protocolo, consulta do dia ou ruptura de estoque virem incidente.",
    ])
    add_heading(doc, "Setor que usa", level=2)
    add_bullets(doc, ["Enfermagem gestora", "Médico", "Dentista", "Técnico de enfermagem", "ACS", "Farmácia", "Gestão"])
    add_heading(doc, "Impacto operacional", level=2)
    add_bullets(doc, [
        "Redução de tempo de briefing de abertura de turno.",
        "Priorização de pacientes em risco sem exigir busca manual.",
        "Visibilidade imediata de cobertura assistencial e gargalos.",
        "Apoio à meta de demanda programada para coordenação de equipe.",
    ])
    add_heading(doc, "Contexto UBS", level=2)
    add_bullets(doc, [
        "Tela é ponto de entrada padrão para quase todos usuários autenticados.",
        "`receptionist` entra em `ReceptionistApp`; `gestor` tende a abrir tab `gestor`.",
        "Dados chegam majoritariamente por bootstrap de sessão.",
    ])

    add_heading(doc, "2. Perfis com Acesso")
    add_table(
        doc,
        ["Perfil", "Acesso", "Escopo", "Restrições", "Observações"],
        [
            ("nurse_manager", "Leitura", "Equipe/unidade", "Sem escrita no painel", "Vê card de demanda"),
            ("doctor", "Leitura", "Equipe/unidade", "Sem escrita no painel", "Pode navegar para pacientes"),
            ("dentist", "Leitura", "Equipe/unidade", "Sem escrita no painel", "Sem card exclusivo de demanda"),
            ("gestor", "Leitura", "Unidade", "Fluxo tende a abrir gestor", "Capability existe"),
            ("acs", "Leitura", "Escopo restrito", "Sem cross-team", "Visão resumida"),
            ("receptionist", "Capability existe", "n/a no fluxo padrão", "Usa ReceptionistApp", "Gap produto/código"),
            ("break_glass_admin", "Leitura", "Global temporário", "Sessão auditada", "Uso excepcional"),
        ],
        [Cm(2.5), Cm(2.5), Cm(2.8), Cm(4.0), Cm(4.6)],
    )
    add_bullets(doc, [
        "Capability base: `dashboard.read`.",
        "Dados do painel respeitam escopo do bootstrap.",
        "Navegação para detalhe continua sujeita às permissões do módulo de destino.",
    ])

    add_heading(doc, "3. Fluxo Principal do Usuário")
    add_numbered(doc, [
        "Usuário autenticado entra no sistema.",
        "Frontend executa `bootstrap(token)`.",
        "Backend retorna `user`, `patients`, `users`, `tasks`, `protocolTemplates` e, quando aplicável, `demandMonthly` e `dataQuality`.",
        "Frontend busca summaries de protocolo por paciente e monta `protocolByPatient`.",
        "Dashboard renderiza KPIs de pacientes, ACS definido, protocolos e profissionais.",
        "Se perfil atual for `nurse_manager`, tela exibe card executivo de demanda programada.",
        "Sistema calcula alertas proativos com base em gravidez em atraso, protocolo crítico, estoque e agenda do dia.",
        "Usuário consulta pacientes prioritários e pode clicar em alertas ou pacientes para navegar.",
        "Usuário usa ações rápidas `Abrir gestão à vista` ou `Ir para pacientes`.",
    ])

    add_heading(doc, "4. Fluxos Alternativos")
    for title, bullets in [
        ("4.1 Bootstrap falha", ["Sistema deve exibir erro claro e impedir falsa leitura operacional."]),
        ("4.2 Sem dados de demanda mensal", ["Card mostra `—`, estado `Sem dados` e explicação de ausência de atendimentos registrados."]),
        ("4.3 Sem pacientes críticos", ["Lista prioritária mostra estado vazio explícito."]),
        ("4.4 Sem alertas proativos", ["Seção de alertas não renderiza; ausência não deve parecer erro."]),
        ("4.5 Sessão expirada / degraded / offline", ["Reautenticar antes de nova carga; painel pode ficar parcial e deve comunicar limitação."]),
    ]:
        add_heading(doc, title, level=2)
        add_bullets(doc, bullets)

    add_heading(doc, "5. Componentes da Tela")
    add_table(
        doc,
        ["Componente", "Tipo", "Obrigatório", "Comportamento", "Observações"],
        [
            ("PageHeader", "Cabeçalho", "Sim", "Título, subtítulo e ações rápidas", "Contexto institucional"),
            ("KPIs", "Cards", "Sim", "Exibem contagens principais", "Leitura imediata"),
            ("Card demanda programada", "Card executivo", "Condicional", "Só para nurse_manager", "Badge + medidor"),
            ("Alertas proativos", "Grid de cards", "Condicional", "Mostra alertas priorizados", "Clicável quando há patientId"),
            ("Pacientes prioritários", "Lista interativa", "Sim", "Mostra até 6 críticos", "Navega ao detalhe"),
            ("Equipe ativa", "Lista", "Sim", "Mostra membros da equipe", "Até 6 itens"),
        ],
        [Cm(3.2), Cm(2.6), Cm(2.1), Cm(5.0), Cm(3.0)],
    )

    add_heading(doc, "6. Campos")
    add_table(
        doc,
        ["Campo", "Tipo", "Obrigatório", "Validação", "Máscara", "Origem", "Observações"],
        [
            ("patients.length", "número", "Sim", "inteiro >= 0", "n/a", "bootstrap", "Pacientes ativos no escopo"),
            ("assignedAcsId", "referência", "Condicional", "vínculo válido", "n/a", "pacientes", "Usado no KPI ACS definido"),
            ("protocolByPatient", "mapa", "Sim", "resumo por patientId", "n/a", "API+cálculo", "Base da criticidade"),
            ("demandMonthly", "objeto", "Condicional", "estrutura válida", "n/a", "bootstrap", "Só perfis autorizados"),
            ("agenda[].date", "data", "Condicional", "data válida", "n/a", "agenda", "Consulta hoje"),
            ("pharmacyStock[].qty", "número", "Condicional", "inteiro >= 0", "n/a", "farmácia", "Risco de estoque"),
            ("currentUser.role", "enum", "Sim", "role válida", "n/a", "sessão", "Altera renderização"),
        ],
        [Cm(2.7), Cm(1.9), Cm(1.9), Cm(4.0), Cm(1.4), Cm(2.0), Cm(4.5)],
    )

    add_heading(doc, "7. Regras de Negócio")
    for rule in [
        ("RN-DASH-001", "Todo usuário com `dashboard.read` pode acessar visão geral do painel.", "Tela de entrada ampla para operação.", "`receptionist` vai para fluxo próprio na UI real."),
        ("RN-DASH-002", "Dados do painel devem respeitar escopo do usuário definido no bootstrap.", "KPIs, listas e alertas não podem expor dados fora do escopo.", "Sessão `break_glass` auditada."),
        ("RN-DASH-003", "Card de demanda programada só aparece para `nurse_manager`.", "Customização por perfil e foco operacional.", "Nenhuma."),
        ("RN-DASH-004", "`demandMonthly` só vem do backend quando usuário é `nurse_manager` ou `doctor`.", "UI precisa suportar ausência do dado.", "Nenhuma."),
        ("RN-DASH-005", "Criticidade de protocolo deve ser calculada por `protocolChip(protocolByPatient[patientId])`.", "KPIs e lista prioritária derivam da mesma regra.", "Fallback seguro sem summary."),
        ("RN-DASH-006", "Alertas proativos devem combinar gestação em atraso, protocolo crítico, estoque crítico e consultas do dia.", "Painel atua como radar institucional.", "Alertas ausentes quando não houver dado."),
        ("RN-DASH-007", "Pacientes prioritários exibem no máximo 6 registros críticos.", "Preserva leitura rápida.", "Nenhuma."),
        ("RN-DASH-008", "Alertas clicáveis navegam apenas quando existir `patientId`.", "Evita ação quebrada para estoque e agenda sem detalhe direto.", "Nenhuma."),
    ]:
        add_rule(doc, *rule)

    add_heading(doc, "8. Auditoria e Logs")
    add_bullets(doc, [
        "Auditar leitura de bootstrap do painel.",
        "Evento real identificado: `admin.bootstrap_read`.",
        "Eventos recomendados: `dashboard.alert_clicked`, `dashboard.patient_priority_opened`, `dashboard.quick_action_triggered`.",
        "Nunca logar CPF completo, CNS completo, endereço completo, tokens ou TOTP.",
    ])

    add_heading(doc, "9. Integrações")
    add_table(
        doc,
        ["Sistema", "Objetivo", "Tipo", "Dados trafegados", "Observações"],
        [
            ("GET /bootstrap", "Carregar contexto inicial", "REST", "user, patients, users, tasks, templates, demanda, qualidade", "Principal integração"),
            ("getProtocolSummaries()", "Montar mapa de protocolos", "REST", "summaries por paciente", "Base de criticidade"),
            ("Agenda", "Alertas do dia", "contexto", "consultas agendadas", "Usada em buildProactiveAlerts"),
            ("Farmácia", "Alertas de estoque", "contexto", "qty, minQty, nome", "Usada em buildProactiveAlerts"),
        ],
        [Cm(2.8), Cm(3.1), Cm(1.7), Cm(5.3), Cm(3.1)],
    )

    add_heading(doc, "10. Estados da Tela")
    add_bullets(doc, [
        "Loading: depende de bootstrap inicial.",
        "Vazio: pacientes prioritários exibem estado vazio claro; alertas podem nem renderizar.",
        "Erro: falha de bootstrap ou summaries deve ser comunicada claramente.",
        "Sucesso: KPIs, alertas e listas refletem contexto carregado da sessão.",
        "Degraded/offline: painel pode ficar parcial e precisa sinalizar limitação.",
    ])

    add_heading(doc, "11. Critérios de Aceite")
    scenarios = [
        ("Cenário 1 — Carregamento padrão do painel", [
            "Dado que usuário autenticado possui `dashboard.read`",
            "Quando acessa sistema",
            "Então painel deve carregar KPIs, listas e ações rápidas compatíveis com seu escopo e backend deve registrar `admin.bootstrap_read`",
        ]),
        ("Cenário 2 — Card de demanda para enfermeira gestora", [
            "Dado que usuária é `nurse_manager`",
            "Quando painel renderiza",
            "Então card de demanda programada deve aparecer e classificar percentual conforme faixa operacional definida",
        ]),
        ("Cenário 3 — Paciente crítico no protocolo", [
            "Dado que paciente possui `protocolChip(...).tone = danger`",
            "Quando painel renderiza lista prioritária",
            "Então paciente deve aparecer entre prioritários e clique deve permitir navegação ao detalhe permitido",
        ]),
        ("Cenário 4 — Receptionist no fluxo real", [
            "Dado que usuário é `receptionist`",
            "Quando autentica no app",
            "Então sistema deve abrir `ReceptionistApp` e não usar painel principal como fluxo padrão",
        ]),
    ]
    for title, lines in scenarios:
        add_heading(doc, title, level=2)
        add_bullets(doc, lines)

    add_heading(doc, "12. Riscos Operacionais")
    add_bullets(doc, [
        "Coordenação confiar em painel desatualizado sem perceber falha de carga.",
        "Bootstrap concentra muitas responsabilidades em uma chamada.",
        "Falha em summaries pode distorcer criticidade do painel.",
        "Equipe deixar de agir por excesso ou falta de alertas.",
        "Diferença entre `dashboard` e `gestor` gerar leituras divergentes.",
    ])

    add_heading(doc, "13. Métricas e Observabilidade")
    add_bullets(doc, [
        "Tempo até primeira ação a partir do painel.",
        "Taxa de clique em alertas proativos.",
        "Taxa de clique em pacientes prioritários.",
        "Latência de `GET /bootstrap` e `getProtocolSummaries`.",
        "Taxa de falha do bootstrap.",
    ])

    add_heading(doc, "14. Dependências Técnicas")
    add_bullets(doc, [
        "Endpoints: `GET /bootstrap` e `getProtocolSummaries(patientIds)`.",
        "Componentes: `Dashboard`, `PageLayout`, `PageHeader`, `KPI`, `Card`, `Chip`, `Badge`, `Avatar`, `Button`.",
        "Permissões: `dashboard.read` + permissões do módulo de destino.",
        "Infraestrutura: autenticação, bootstrap, summaries, agenda e farmácia no contexto do app.",
    ])

    add_heading(doc, "15. Considerações UX Operacionais")
    add_bullets(doc, [
        "Painel deve funcionar como radar do turno.",
        "Ações rápidas devem reduzir busca manual por módulo.",
        "Badges e chips precisam combinar cor + texto.",
        "Alertas proativos devem priorizar relevância real e evitar ruído excessivo.",
    ])

    add_heading(doc, "16. Gaps Identificados")
    add_bullets(doc, [
        "`allUsers` é passado ao componente, mas não é usado no `Dashboard` atual.",
        "`dataQuality` é retornado no bootstrap, mas não é explorado no painel atual.",
        "`receptionist` possui `dashboard.read` em RBAC, mas fluxo real ignora painel principal.",
        "Painel depende de `getProtocolSummaries` extra após bootstrap; oportunidade de consolidação.",
        "Alertas de estoque e agenda nem sempre possuem destino navegável direto.",
    ])

    add_heading(doc, "17. Melhorias Recomendadas")
    add_bullets(doc, [
        "Incorporar `dataQuality` no painel.",
        "Explicitar timestamp de última carga.",
        "Mostrar badge de degradado/desatualizado em carga parcial.",
        "Revisar necessidade de `dashboard.read` para `receptionist`.",
        "Consolidar summaries de protocolo no próprio bootstrap.",
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
