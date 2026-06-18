from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\dev\vitras")
OUTPUT = ROOT / "docs" / "escopos" / "fila-recepcao.docx"


ACCENT = RGBColor(14, 77, 146)
ACCENT_DARK = RGBColor(24, 45, 74)
TEXT = RGBColor(40, 40, 40)
MUTED = RGBColor(92, 104, 117)
LIGHT_FILL = "EAF2FB"
HEADER_FILL = "D7E6F7"
BORDER = "C7D4E3"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=TEXT, size=10.5, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, widths=None):
    table.style = "Table Grid"
    table.autofit = False
    if widths:
      for row in table.rows:
          for idx, width in enumerate(widths):
              if idx < len(row.cells):
                  row.cells[idx].width = width
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "left", "bottom", "right"):
                element = tc_mar.find(qn(f"w:{side}"))
                if element is None:
                    element = OxmlElement(f"w:{side}")
                    tc_mar.append(element)
                element.set(qn("w:w"), "110")
                element.set(qn("w:type"), "dxa")
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, HEADER_FILL)
            set_repeat_table_header(row)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT_DARK
    p.space_after = Pt(6)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = "Aptos"
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    p.space_after = Pt(12)


def add_meta_table(doc):
    table = doc.add_table(rows=4, cols=2)
    style_table(table, [Cm(4.2), Cm(11.8)])
    meta = [
        ("Produto", "VITRAS"),
        ("Módulo", "Recepção / Fila de Atendimento / Triagem inicial / Agenda operacional do dia"),
        ("Tab/Rota", "`queue` no shell principal + app dedicada `ReceptionistApp` para `receptionist`"),
        ("Última atualização", "27/05/2026"),
    ]
    for i, (label, value) in enumerate(meta):
        set_cell_text(table.cell(i, 0), label, bold=True, color=ACCENT_DARK)
        set_cell_text(table.cell(i, 1), value)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    r.font.color.rgb = ACCENT_DARK
    r.font.name = "Aptos Display"
    if level == 1:
        r.font.size = Pt(15)
    elif level == 2:
        r.font.size = Pt(12.5)
    else:
        r.font.size = Pt(11.5)
    p.space_before = Pt(8)
    p.space_after = Pt(4)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(item)
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(item)
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = ACCENT_DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = TEXT
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, color=ACCENT_DARK, align="center")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_rule_block(doc, rule_id, description, impact, exception):
    add_paragraph(doc, description, bold_prefix=f"{rule_id} — Descrição: ")
    add_paragraph(doc, impact, bold_prefix="Impacto: ")
    add_paragraph(doc, exception, bold_prefix="Exceção: ")


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "Fila / Recepção")
    add_subtitle(doc, "Especificação funcional enterprise para operação UBS, produto, QA e engenharia.")
    add_meta_table(doc)

    add_heading(doc, "1. Objetivo da Tela")
    add_heading(doc, "Qual problema resolve", level=2)
    add_bullets(doc, [
        "Centraliza chegada de pacientes do dia em fluxo único, ordenado e auditável.",
        "Reduz uso informal de papel e recados verbais entre recepção, triagem e assistência.",
        "Diminui erro operacional de chamar paciente fora de prioridade clínica ou social.",
        "Dá visibilidade imediata de espera, atendimento em curso e volume diário.",
    ])
    add_heading(doc, "Setor que usa", level=2)
    add_bullets(doc, ["Recepção", "Técnico de enfermagem", "Enfermagem gestora", "Apoio operacional da UBS"])
    add_heading(doc, "Impacto operacional", level=2)
    add_bullets(doc, [
        "Reduz tempo entre chegada e encaminhamento para atendimento.",
        "Padroniza classificação inicial de prioridade de recepção.",
        "Evita duplicidade de entrada do mesmo paciente na fila ativa.",
        "Melhora rastreabilidade de quem colocou, chamou, concluiu ou removeu paciente da fila.",
    ])
    add_heading(doc, "Contexto UBS", level=2)
    add_bullets(doc, [
        "Usada na chegada do paciente à unidade e durante todo turno de atendimento.",
        "Recepção usa visão simplificada em `ReceptionistApp`; equipe técnica usa `QueuePage` com mais contexto clínico-operacional.",
        "Fluxo depende de paciente previamente cadastrado ou cadastro rápido durante entrada.",
        "Fila pertence à equipe do usuário autenticado. Sem `teamId`, acesso deve falhar por segurança.",
    ])

    add_heading(doc, "2. Perfis com Acesso")
    add_table(
        doc,
        ["Perfil", "Acesso", "Escopo", "Restrições", "Observações"],
        [
            ("receptionist", "Leitura e escrita", "Própria equipe", "Sem escrita clínica; sem cross-team", "Usa app dedicada"),
            ("nursing_tech", "Leitura e escrita", "Própria equipe", "Fluxo de triagem posterior", "Principal operador clínico"),
            ("nurse_manager", "Sem capability direta hoje", "n/a", "Não acessa queue.read/write no código atual", "Gap de produto/código"),
            ("doctor", "Sem acesso direto", "n/a", "Recebe paciente via atendimento", "Sem queue.read/write"),
            ("gestor", "Sem acesso direto", "n/a", "Sem uso operacional", "Visão indireta por indicadores"),
            ("break_glass_admin", "Leitura e escrita", "Global emergencial", "Acesso temporário auditado", "Uso excepcional"),
        ],
        [Cm(2.5), Cm(3.0), Cm(2.8), Cm(4.1), Cm(4.0)],
    )
    add_bullets(doc, [
        "Leitura exige `queue.read` ou `queue.write`.",
        "Escrita exige `queue.write`.",
        "Filtro backend: `entry.teamId === req.user.teamId`.",
        "Paciente só entra na fila se `patient.teamId === req.user.teamId`.",
    ])

    add_heading(doc, "3. Fluxo Principal do Usuário")
    add_numbered(doc, [
        "Usuário autenticado acessa módulo de fila.",
        "Sistema valida sessão, capability e `teamId`.",
        "Frontend carrega fila via `GET /queue` e inicia polling a cada 15 segundos.",
        "Usuário visualiza KPIs de aguardando, em atendimento e total do dia.",
        "Usuário aciona `Dar entrada` e busca paciente existente por nome, CPF ou telefone.",
        "Se paciente não existir no shell principal `QueuePage`, operador pode cadastrar novo paciente e retornar ao fluxo.",
        "Sistema sugere prioridade automática por categoria do cuidado e idade.",
        "Usuário informa tipo de demanda (`scheduled` ou `spontaneous`) e, se espontânea, define destino (`doctor` ou `nurse`).",
        "Sistema valida duplicidade na fila ativa e cria entrada com status inicial `waiting`.",
        "Operador avança status conforme atendimento: `waiting -> triage -> ready -> attending -> done`.",
        "Ao fim do ciclo, operador pode limpar concluídos via `POST /queue/clear-done`.",
    ])

    add_heading(doc, "4. Fluxos Alternativos")
    for title, bullets in [
        ("4.1 Paciente não encontrado", [
            "Em `QueuePage`, operador pode cadastrar paciente novo sem sair do fluxo.",
            "Em `ReceptionistApp`, fluxo atual depende de paciente já existente.",
            "Busca sem retorno deve comunicar ausência de cadastro com clareza.",
        ]),
        ("4.2 Paciente já está na fila ativa", [
            "Backend retorna `409` com mensagem `Paciente já está na fila ativa`.",
            "Frontend deve bloquear nova entrada silenciosa e orientar uso da entrada já existente.",
        ]),
        ("4.3 Paciente fora da equipe atual", [
            "Backend retorna `403` e bloqueia enfileiramento cross-team.",
            "Caso legítimo exige fluxo separado e auditado; não usar fila comum.",
        ]),
        ("4.4 Concorrência operacional", [
            "Código atual não usa controle explícito de versão na fila.",
            "Risco maior em campos `reason`, `priority`, `vitals`, `triage*`.",
        ]),
        ("4.5 Degraded / offline / sessão expirada", [
            "Leitura pode permanecer parcial; escrita deve falhar com mensagem clara.",
            "Polling não deve travar UI.",
            "Sessão expirada deve exigir novo login antes de persistir qualquer ação.",
        ]),
    ]:
        add_heading(doc, title, level=2)
        add_bullets(doc, bullets)

    add_heading(doc, "5. Componentes da Tela")
    add_table(
        doc,
        ["Componente", "Tipo", "Obrigatório", "Comportamento", "Observações"],
        [
            ("Cabeçalho", "Header", "Sim", "Exibe contexto e CTA principal", "Título varia por shell"),
            ("KPIs", "Cards", "Sim", "Mostram waiting, attending, total", "Atualização por polling"),
            ("Lista da fila", "Lista ordenada", "Sim", "Renderiza posição, badges e ações", "Prioridade + chegada"),
            ("Modal Dar entrada", "Modal", "Sim", "Cria entrada na fila", "Busca avançada em QueuePage"),
            ("Cadastro rápido", "Form modal", "Condicional", "Cadastra paciente novo", "Só em QueuePage"),
            ("Botão Chamar", "Ação", "Condicional", "Avança status", "Diferença entre shells"),
            ("Botão Limpar concluídos", "Ação coletiva", "Condicional", "Marca done como cleared", "Só se houver concluídos"),
        ],
        [Cm(3.0), Cm(2.4), Cm(2.1), Cm(5.0), Cm(3.2)],
    )

    add_heading(doc, "6. Campos")
    add_table(
        doc,
        ["Campo", "Tipo", "Obrigatório", "Validação", "Máscara", "Origem", "Observações"],
        [
            ("patientId", "referência", "Sim", "UUID válido; mesma equipe", "n/a", "usuário/API", "Obrigatório em criação"),
            ("priority", "enum", "Sim", "urgent, elderly, pregnant, child, normal", "n/a", "usuário/sistema", "Pode ser inferido"),
            ("reason", "texto curto", "Não", "até 1000 caracteres", "n/a", "usuário", "Queixa principal operacional"),
            ("demandType", "enum", "Não", "scheduled ou spontaneous", "n/a", "usuário", "Default scheduled"),
            ("destination", "enum", "Condicional", "doctor ou nurse", "n/a", "usuário", "Obrigatório na UX para espontânea"),
            ("status", "enum", "Automático", "waiting, triage, ready, attending, done", "n/a", "sistema", "removed e cleared fora do patch schema"),
            ("arrivedAt", "datetime", "Automático", "ISO válido", "n/a", "backend", "Ordenação e tempo de espera"),
        ],
        [Cm(2.4), Cm(1.9), Cm(1.8), Cm(4.6), Cm(1.5), Cm(2.0), Cm(4.2)],
    )

    add_heading(doc, "7. Regras de Negócio")
    rules = [
        ("RN-QUEUE-001", "Somente usuários com `queue.read` ou `queue.write` podem visualizar fila.", "Controle de acesso de tela e API.", "Nenhuma."),
        ("RN-QUEUE-002", "Somente usuários com `queue.write` podem criar, atualizar, remover ou limpar fila.", "Ações de mutação bloqueadas para leitura apenas.", "Nenhuma."),
        ("RN-QUEUE-003", "Fila é sempre filtrada por `teamId` do usuário autenticado.", "Isolamento multi-tenant intraunidade/equipe.", "`team.manage` em operação administrativa elevada."),
        ("RN-QUEUE-004", "Paciente só pode entrar na fila se pertencer à mesma equipe do usuário.", "Evita mistura de fluxo operacional entre equipes.", "Não implementada para cross-team comum."),
        ("RN-QUEUE-005", "Paciente não pode possuir duas entradas simultâneas na fila ativa.", "Backend rejeita duplicidade com `409`.", "Nenhuma."),
        ("RN-QUEUE-006", "Ordenação da fila deve obedecer prioridade e, em empate, horário de chegada.", "Lista exibida e percepção operacional.", "Itens `done` vão para fim."),
        ("RN-QUEUE-007", "Remoção da fila não exclui fisicamente registro; troca status para `removed`.", "Rastreabilidade e auditoria preservadas.", "Nenhuma."),
        ("RN-QUEUE-008", "Limpeza de concluídos afeta apenas entradas `done` da equipe atual, convertendo-as para `cleared`.", "Encerramento operacional sem perda de histórico.", "Nenhuma."),
    ]
    for rule in rules:
        add_rule_block(doc, *rule)

    add_heading(doc, "8. Auditoria e Logs")
    add_bullets(doc, [
        "Auditar criação, atualização de status, mudança de prioridade, remoção e limpeza de concluídos.",
        "Eventos reais: `queue.entry_created`, `queue.entry_updated`, `queue.entry_removed`, `queue.done_cleared`.",
        "Campos mínimos: `patientId`, `teamId`, `before`, `after`, `changedFields`, `actor`, `request context`.",
        "Nunca logar CPF completo, CNS completo, telefone completo, endereço completo, tokens ou TOTP.",
    ])

    add_heading(doc, "9. Integrações")
    add_table(
        doc,
        ["Sistema", "Objetivo", "Tipo", "Dados trafegados", "Observações"],
        [
            ("API `/queue`", "CRUD da fila", "REST", "paciente, prioridade, motivo, demanda, status", "Principal integração"),
            ("API `/patients`", "Busca/cadastro rápido", "REST", "identificação e dados cadastrais", "Usada em QueuePage"),
            ("API `/agenda`", "Contexto de chegada programada", "REST", "data, hora, profissional, tipo", "Sem vínculo forte ainda"),
            ("Auditoria interna", "Rastreabilidade", "serviço", "snapshots before/after", "Obrigatória"),
            ("Redis/rate limit", "Proteção operacional", "infra", "contadores de requisição", "Fail-closed em produção"),
        ],
        [Cm(2.6), Cm(3.0), Cm(1.7), Cm(5.3), Cm(3.0)],
    )

    add_heading(doc, "10. Estados da Tela")
    add_bullets(doc, [
        "Loading: `useQueue` carrega lista e faz polling a cada 15s.",
        "Vazio: mostrar `Nenhum paciente na fila` com CTA `Dar entrada`.",
        "Erro: banner claro sem perder contexto visual.",
        "Sucesso: atualizar lista, fechar modal e limpar formulário.",
        "Degraded: bloquear escrita se backend não estiver confiável.",
        "Sem permissão: negar acesso sem expor regra interna detalhada.",
        "Offline parcial: bloquear mutações até retorno de conectividade.",
    ])

    add_heading(doc, "11. Critérios de Aceite")
    scenarios = [
        ("Cenário 1 — Entrada de paciente agendado", [
            "Dado que um `receptionist` autenticado possui `queue.write` e paciente pertence à mesma equipe",
            "Quando o operador der entrada com `demandType = scheduled`",
            "Então o sistema deve criar entrada com status `waiting`, ordenar paciente corretamente e registrar `queue.entry_created`",
        ]),
        ("Cenário 2 — Bloqueio de duplicidade", [
            "Dado que o paciente já possui entrada com status ativo",
            "Quando o operador tentar criar nova entrada",
            "Então o backend deve retornar `409` e interface deve informar que paciente já está na fila ativa",
        ]),
        ("Cenário 3 — Isolamento por equipe", [
            "Dado que paciente pertence a equipe diferente da equipe do operador",
            "Quando o operador tentar enfileirá-lo",
            "Então o sistema deve negar operação com `403` e não criar entrada",
        ]),
        ("Cenário 4 — Conclusão do atendimento", [
            "Dado que entrada está em `attending`",
            "Quando operador clicar em `Concluir`",
            "Então sistema deve mudar status para `done` e manter item elegível para `Limpar concluídos`",
        ]),
    ]
    for title, lines in scenarios:
        add_heading(doc, title, level=2)
        add_bullets(doc, lines)

    add_heading(doc, "12. Riscos Operacionais")
    add_bullets(doc, [
        "Recepção chamar paciente errado por nomes parecidos.",
        "Atalho `waiting -> attending` na app simplificada pode pular triagem obrigatória.",
        "Polling de 15s pode deixar percepção de fila desatualizada por alguns segundos.",
        "Schema de `vitals` permissivo demais.",
        "Cadastro rápido em recepção coleta dados sensíveis demais para alguns cenários de balcão.",
    ])

    add_heading(doc, "13. Métricas e Observabilidade")
    add_bullets(doc, [
        "Tempo médio entre `arrivedAt` e `attending`.",
        "Tempo médio entre `arrivedAt` e `done`.",
        "Quantidade diária de pacientes por prioridade.",
        "Proporção `scheduled` vs `spontaneous`.",
        "Taxa de `409` em criação e pico de `403` por unidade.",
        "Alertar alto uso de `break_glass_admin` em recepção.",
    ])

    add_heading(doc, "14. Dependências Técnicas")
    add_bullets(doc, [
        "Endpoints: `GET /queue`, `POST /queue`, `PATCH /queue/:id`, `DELETE /queue/:id`, `POST /queue/clear-done`.",
        "Hooks/componentes: `useQueue`, `QueuePage`, `ReceptionistApp`, `Modal`, `Input`, `Select`, `Button`, `KPI`.",
        "Permissões: `queue.read`, `queue.write`, `team.manage` em casos elevados.",
        "Infraestrutura: autenticação, PostgreSQL ou file-mode em dev, Redis/rate limit, auditoria ativa.",
    ])

    add_heading(doc, "15. Considerações UX Operacionais")
    add_bullets(doc, [
        "Recepção precisa registrar chegada em poucos segundos.",
        "Lista deve ser legível sob pressão de atendimento.",
        "Sugestão automática de prioridade reduz cliques.",
        "Badges devem combinar texto + cor.",
        "Ideal futuro: confirmação antes de remover entrada.",
    ])

    add_heading(doc, "16. Gaps Identificados")
    add_bullets(doc, [
        "`nurse_manager` não possui acesso de fila no RBAC atual.",
        "`ReceptionistApp` pula `triage/ready` e chama direto `attending`.",
        "Controle de concorrência é fraco para mutações simultâneas.",
        "`vitals` aceita estrutura aberta no backend.",
        "`agendaRef` existe no schema, mas está pouco explorado no fluxo.",
    ])

    add_heading(doc, "17. Melhorias Recomendadas")
    add_bullets(doc, [
        "Criar modo explícito `recepção pura` vs `triagem`.",
        "Adicionar confirmação com motivo para remoção da fila.",
        "Incluir lock otimista por `updatedAt`.",
        "Reforçar vínculo entre entrada agendada e `agendaRef`.",
        "Avaliar atualização near-real-time por SSE/WebSocket se volume crescer.",
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
